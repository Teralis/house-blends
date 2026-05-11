"""
NSW Ports Weekly Hold Point Report — Python-only generator.

CLI usage:  python src/generate_weekly_report.py [csv_file_path]
            Output written to output/NSW_Ports_Weekly_Report_YYYYMMDD.docx

Streamlit:  import analyze_submittals, generate_word_report directly.
            generate_word_report() returns a BytesIO ready for st.download_button.
"""

import io
import os
import sys
import warnings
from datetime import datetime, timedelta

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

warnings.filterwarnings('ignore')

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'output')


# ---------------------------------------------------------------------------
# Data analysis
# ---------------------------------------------------------------------------

def get_week_boundaries(reference_date=None):
    if reference_date is None:
        reference_date = datetime.now()
    days_since_sunday = (reference_date.weekday() + 1) % 7
    last_sunday = reference_date - timedelta(days=days_since_sunday)
    last_sunday = last_sunday.replace(hour=23, minute=59, second=59)
    week_start = last_sunday - timedelta(days=6)
    week_start = week_start.replace(hour=0, minute=0, second=0)
    return week_start, last_sunday


def analyze_submittals(csv_source):
    """
    Accept a file path or file-like object (e.g. Streamlit UploadedFile).
    Returns a report_data dict consumed by generate_word_report().
    """
    try:
        df = pd.read_csv(csv_source, encoding='utf-8-sig')
    except Exception:
        df = pd.read_csv(csv_source, encoding='latin-1')

    df.columns = df.columns.str.lower()

    if '#' not in df.columns and 'submittal number' in df.columns:
        df = df.rename(columns={'submittal number': '#'})

    date_columns = ['created at', 'updated at', 'sent date', 'returned date',
                    'wf due date', 'final due date', 'distributed date']
    for col in date_columns:
        if col in df.columns:
            df[col] = pd.to_datetime(df[col], format='%d/%m/%Y', errors='coerce')
            if df[col].isna().all():
                df[col] = pd.to_datetime(df[col], format='%d/%m/%Y at %I:%M %p', errors='coerce')

    nsw_ports_mask = (
        df['approvers'].str.contains('NSW Ports', na=False, case=False) |
        df['action required by'].str.contains('NSW Ports', na=False, case=False)
    )
    nsw_ports_df = df[nsw_ports_mask].copy()

    hold_points_df = nsw_ports_df[
        nsw_ports_df['type'].str.contains('Hold Point', na=False, case=False)
    ].copy()

    hold_points_df['Base_HPWP'] = hold_points_df['#'].astype(str)
    hold_points_df['rev.'] = pd.to_numeric(hold_points_df['rev.'], errors='coerce').fillna(0).astype(int)
    hold_points_df['HPWP_Display'] = hold_points_df.apply(
        lambda r: f"{r['#']} Rev {r['rev.']}" if r['rev.'] > 0 else r['#'], axis=1
    )

    latest_revisions = (
        hold_points_df.sort_values('rev.', ascending=False)
        .groupby('Base_HPWP').first().reset_index()
    )

    week_start, week_end = get_week_boundaries()

    week_activity = hold_points_df[
        ((hold_points_df['created at']   >= week_start) & (hold_points_df['created at']   <= week_end)) |
        ((hold_points_df['sent date']    >= week_start) & (hold_points_df['sent date']    <= week_end)) |
        ((hold_points_df['returned date'] >= week_start) & (hold_points_df['returned date'] <= week_end)) |
        ((hold_points_df['updated at']   >= week_start) & (hold_points_df['updated at']   <= week_end))
    ].copy()

    new_submissions  = hold_points_df[
        (hold_points_df['sent date'] >= week_start) & (hold_points_df['sent date'] <= week_end)
    ].copy()

    returns_responses = hold_points_df[
        (hold_points_df['returned date'] >= week_start) & (hold_points_df['returned date'] <= week_end)
    ].copy()

    return {
        'week_start':        week_start,
        'week_end':          week_end,
        'total_hold_points': latest_revisions['Base_HPWP'].nunique(),
        'week_activity_count': week_activity['Base_HPWP'].nunique(),
        'status_counts':     latest_revisions['status'].value_counts().to_dict(),
        'response_counts':   latest_revisions['response'].value_counts().to_dict(),
        'hold_points':       latest_revisions,
        'week_activity':     week_activity,
        'new_submissions':   new_submissions,
        'returns_responses': returns_responses,
    }


# ---------------------------------------------------------------------------
# Word generation helpers
# ---------------------------------------------------------------------------

def _status_color(response):
    r = str(response)
    if 'Not Released'        in r: return RGBColor(0xCC, 0x00, 0x00)
    if 'Released with Cond'  in r: return RGBColor(0xFF, 0x88, 0x00)
    if 'Released'            in r: return RGBColor(0x00, 0xAA, 0x00)
    if 'Pending'             in r: return RGBColor(0x00, 0x66, 0xCC)
    return RGBColor(0x00, 0x00, 0x00)


def _set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:color'), 'CCCCCC')
        tcPr.append(el)


def _set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)


def _set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    trPr.append(el)


def _bullet(doc, text, color=None, bold=False, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if bold:  run.font.bold = True
    return p


# ---------------------------------------------------------------------------
# Word generation
# ---------------------------------------------------------------------------

def generate_word_report(report_data):
    """Return a BytesIO containing the .docx report."""
    week_start = report_data['week_start']
    week_end   = report_data['week_end']

    doc = Document()

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    # Title
    title = doc.add_heading('NSW Ports Hold Point Weekly Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Week Ending: {week_end.strftime('%A, %d %B %Y')}")
    run.italic = True
    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        f"This report provides an overview of Hold Point activity for NSW Ports "
        f"for the week ending {week_end.strftime('%A, %d %B %Y')}. "
        f"The report focuses on Hold Points requiring NSW Ports approval or action."
    )
    doc.add_paragraph()

    # Key Metrics
    doc.add_heading('Key Metrics', level=2)
    open_count   = report_data['status_counts'].get('Open',   0)
    closed_count = report_data['status_counts'].get('Closed', 0)
    _bullet(doc, f"Total NSW Ports Hold Points tracked: {report_data['total_hold_points']}")
    _bullet(doc, f"Hold Points with activity this week: {report_data['week_activity_count']}")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Current Status: ").font.size = Pt(11)
    r = p.add_run(f"{open_count} Open");   r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xFF, 0x88, 0x00)
    p.add_run(", ").font.size = Pt(11)
    r = p.add_run(f"{closed_count} Closed"); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x00, 0xAA, 0x00)
    doc.add_paragraph()

    # Week's Activity
    doc.add_heading("Week's Activity Summary", level=2)
    doc.add_paragraph(
        f"The following Hold Points had activity during the week "
        f"({week_start.strftime('%d %B %Y')} – {week_end.strftime('%d %B %Y')}):"
    )
    doc.add_paragraph()

    new_subs = report_data['new_submissions'].sort_values('sent date')
    returns  = report_data['returns_responses'].sort_values('returned date')

    if not new_subs.empty:
        doc.add_paragraph().add_run("New Submissions").bold = True
        for _, row in new_subs.iterrows():
            title = str(row['title'])[:60]
            date  = row['sent date'].strftime('%d/%m') if pd.notna(row['sent date']) else ''
            _bullet(doc, f"{row['HPWP_Display']} – {title} (Sent {date})")
        doc.add_paragraph()

    if not returns.empty:
        doc.add_paragraph().add_run("Returns / Responses").bold = True
        for _, row in returns.iterrows():
            title    = str(row['title'])[:50]
            response = row['response'] if pd.notna(row['response']) else 'Pending'
            date     = row['returned date'].strftime('%d/%m') if pd.notna(row['returned date']) else ''
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{row['HPWP_Display']} – {title} – ").font.size = Pt(11)
            r = p.add_run(response.upper())
            r.bold = True; r.font.size = Pt(11); r.font.color.rgb = _status_color(response)
            p.add_run(f" ({date})").font.size = Pt(11)
        doc.add_paragraph()

    if new_subs.empty and returns.empty:
        doc.add_paragraph("No new activity this week.").runs[0].italic = True

    doc.add_page_break()

    # Status table
    doc.add_heading('All NSW Ports Hold Points Status', level=1)
    doc.add_paragraph(f"Status as of {week_end.strftime('%A, %d %B %Y')}:")
    doc.add_paragraph()

    hold_points = report_data['hold_points'].sort_values('#')
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    for i, w in enumerate([Inches(1.1), Inches(2.5), Inches(1.1), Inches(1.6), Inches(1.1)]):
        table.columns[i].width = w

    hdr = table.rows[0]
    _set_repeat_header(hdr)
    for i, (cell, text) in enumerate(zip(hdr.cells, ['HP Number', 'Title', 'Submitted', 'Response', 'Status'])):
        _set_cell_border(cell)
        _set_cell_shading(cell, 'D5E8F0')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text); run.bold = True; run.font.size = Pt(10)

    for _, row in hold_points.iterrows():
        submitted = row['sent date'].strftime('%d/%m/%Y') if pd.notna(row['sent date']) else 'Not Sent'
        response  = row['response'] if pd.notna(row['response']) else 'No Response'
        status    = row['status']
        values    = [row['HPWP_Display'], str(row['title'])[:45], submitted, response, status]
        aligns    = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                     WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
        tr = table.add_row()
        for i, (cell, val, align) in enumerate(zip(tr.cells, values, aligns)):
            _set_cell_border(cell)
            p = cell.paragraphs[0]; p.alignment = align
            run = p.add_run(val); run.font.size = Pt(9)
            if i == 3:
                run.bold = True; run.font.color.rgb = _status_color(val)
            if i == 4:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0x88, 0x00) if val == 'Open' else RGBColor(0x00, 0xAA, 0x00)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    csv_path = sys.argv[1] if len(sys.argv) > 1 else 'SubmittalLog.csv'
    if not os.path.exists(csv_path):
        sys.exit(f"ERROR: {csv_path} not found")

    print(f"Processing: {csv_path}")
    report_data = analyze_submittals(csv_path)
    buf         = generate_word_report(report_data)

    week_end  = report_data['week_end']
    filename  = f"NSW_Ports_Weekly_Report_{week_end.strftime('%Y%m%d')}.docx"
    out_path  = os.path.join(OUTPUT_DIR, filename)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    with open(out_path, 'wb') as f:
        f.write(buf.read())

    print(f"Generated: output/{filename}")


if __name__ == "__main__":
    main()
